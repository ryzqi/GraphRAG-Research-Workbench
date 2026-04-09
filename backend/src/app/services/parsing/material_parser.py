from __future__ import annotations

import asyncio
import os
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

import httpx

from app.core.settings import Settings, get_settings
from app.integrations.object_storage import ObjectRef, ObjectStorage
from app.models.source_material import SourceMaterial, SourceType
from app.services.parsing.errors import ParseError
from app.services.parsing.types import ParsedChunk, ParsedDocument

_UPLOAD_EXT_TO_KIND: dict[str, str] = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".md": "md",
    ".txt": "txt",
    ".doc": "doc",
}

_UPLOAD_MIME_TO_KIND: dict[str, str] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/markdown": "md",
    "text/plain": "txt",
    "application/msword": "doc",
}


def _has_meaningful_content(doc: ParsedDocument) -> bool:
    text_ok = bool(doc.text and doc.text.strip())
    chunk_ok = any(
        bool(chunk.text and chunk.text.strip()) for chunk in (doc.chunks or [])
    )
    return text_ok or chunk_ok


def _ensure_non_empty(doc: ParsedDocument) -> ParsedDocument:
    if not _has_meaningful_content(doc):
        raise ParseError(
            error_code="EMPTY_PARSE_RESULT",
            message="解析结果为空",
        )
    return doc


def _decode_text_bytes(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("gb18030", errors="replace")


def _infer_upload_kind(*, mime_type: str | None, extension: str | None) -> str | None:
    mime_kind = _UPLOAD_MIME_TO_KIND.get(mime_type or "")
    ext_kind = _UPLOAD_EXT_TO_KIND.get(extension or "")

    if mime_kind == "doc" or ext_kind == "doc":
        raise ParseError(
            error_code="DOC_NOT_SUPPORTED",
            message="不支持 .doc（请转换为 .docx 后再导入）",
        )

    if mime_kind and ext_kind and mime_kind != ext_kind:
        raise ParseError(
            error_code="MIME_EXTENSION_MISMATCH",
            message=f"MIME 与扩展名冲突：mime_type={mime_type!r} extension={extension!r}",
            details={"mime_type": mime_type, "extension": extension},
        )

    return mime_kind or ext_kind


def _extract_minio_ref(uri: str) -> tuple[str, str]:
    if not uri.startswith("minio://"):
        raise ParseError(
            error_code="INVALID_URI", message="UPLOAD 资料 uri 不是 minio://"
        )

    uri_parts = uri[8:].split("/", 1)
    if len(uri_parts) != 2:
        raise ParseError(error_code="INVALID_URI", message="UPLOAD 资料 uri 格式非法")

    bucket, object_name = uri_parts
    if not bucket or not object_name:
        raise ParseError(
            error_code="INVALID_URI", message="UPLOAD 资料 uri 缺少 bucket/object_name"
        )
    return bucket, object_name


def _material_filename_from_object_name(object_name: str) -> str:
    # material_service.upload_file 生成的 object_name 形如 f"{kb_id}/{material_id}/{filename}"
    return object_name.split("/")[-1] if object_name else ""


def _page_blocks_from_middle_json(
    page: dict[str, Any],
) -> tuple[str, list[dict[str, Any]]]:
    para_blocks = page.get("para_blocks")
    if isinstance(para_blocks, list) and para_blocks:
        return "para_blocks", [b for b in para_blocks if isinstance(b, dict)]

    # 兜底：当 para_blocks 缺失或为空时，尝试使用 preproc_blocks，避免文本被丢弃。
    preproc_blocks = page.get("preproc_blocks")
    if isinstance(preproc_blocks, list):
        return "preproc_blocks", [b for b in preproc_blocks if isinstance(b, dict)]
    return "none", []


def _collect_chunks_from_mineru_block(
    *,
    block: dict[str, Any],
    page_idx: int,
    block_id: str,
    block_source: str,
    chunks: list[ParsedChunk],
) -> None:
    text = _mineru_block_to_text(block).strip()
    if text:
        block_entry: dict[str, Any] = {"id": block_id}
        if "bbox" in block:
            block_entry["bbox"] = block.get("bbox")
        if "type" in block:
            block_entry["type"] = block.get("type")

        locator: dict[str, Any] = {
            "kind": "pdf",
            "page_start": page_idx,
            "page_end": page_idx,
            "blocks": [block_entry],
        }
        chunks.append(
            ParsedChunk(
                text=text,
                locator=locator,
                metadata={
                    "mineru_block_type": block.get("type"),
                    "mineru_block_id": block_id,
                    "mineru_block_source": block_source,
                },
            )
        )

    child_blocks = block.get("blocks")
    if not isinstance(child_blocks, list):
        return

    for child_index, child in enumerate(child_blocks):
        if not isinstance(child, dict):
            continue
        child_id = f"{block_id}_c{child_index}"
        _collect_chunks_from_mineru_block(
            block=child,
            page_idx=page_idx,
            block_id=child_id,
            block_source=block_source,
            chunks=chunks,
        )


def extract_pdf_chunks_from_middle_json(
    middle_json: dict[str, Any],
) -> list[ParsedChunk]:
    """从 MinerU middle.json 中提取 block=chunk 的切片（用于 PDF locator）。"""
    pdf_info = middle_json.get("pdf_info")
    if not isinstance(pdf_info, list):
        return []

    chunks: list[ParsedChunk] = []
    for page in pdf_info:
        if not isinstance(page, dict):
            continue

        page_idx = page.get("page_idx")
        if not isinstance(page_idx, int):
            continue

        block_source, blocks = _page_blocks_from_middle_json(page)
        for block_index, block in enumerate(blocks):
            block_id = f"p{page_idx}_b{block_index}"
            _collect_chunks_from_mineru_block(
                block=block,
                page_idx=page_idx,
                block_id=block_id,
                block_source=block_source,
                chunks=chunks,
            )

    return chunks


def _mineru_block_to_text(block: dict[str, Any]) -> str:
    # middle.json 的结构为：block -> lines -> spans -> content
    lines = block.get("lines")
    if not isinstance(lines, list):
        return ""

    line_texts: list[str] = []
    for line in lines:
        if not isinstance(line, dict):
            continue
        spans = line.get("spans")
        if not isinstance(spans, list):
            continue
        span_texts: list[str] = []
        for span in spans:
            if not isinstance(span, dict):
                continue
            content = span.get("content")
            if isinstance(content, str):
                text = content.strip()
                if text:
                    span_texts.append(text)
        if span_texts:
            line_texts.append("".join(span_texts))
    return "\n".join(line_texts)


async def parse_material(
    material: SourceMaterial,
    *,
    settings: Settings | None = None,
    http_client: httpx.AsyncClient | None = None,
    storage: ObjectStorage | None = None,
) -> ParsedDocument:
    """将 SourceMaterial 解析为 ParsedDocument（失败抛 ParseError）。"""
    cfg = settings or get_settings()

    if material.source_type == SourceType.TEXT:
        text = (material.metadata_ or {}).get("text")
        if not isinstance(text, str):
            text = ""
        return _ensure_non_empty(
            ParsedDocument(
                text=text,
                mime_type=material.mime_type,
                locator=None,
                metadata={"source_type": material.source_type.value},
            )
        )

    if material.source_type == SourceType.URL:
        if not material.uri:
            raise ParseError(error_code="INVALID_URL", message="URL 资料缺少 uri")
        if http_client is None:
            raise ParseError(
                error_code="MISSING_HTTP_CLIENT", message="URL 解析缺少 http_client"
            )
        doc = await _parse_url(material.uri, http_client=http_client, settings=cfg)
        doc.mime_type = doc.mime_type or material.mime_type
        doc.metadata = {
            **(doc.metadata or {}),
            "source_type": material.source_type.value,
        }
        return _ensure_non_empty(doc)

    if material.source_type == SourceType.UPLOAD:
        if not material.uri:
            raise ParseError(error_code="INVALID_URI", message="UPLOAD 资料缺少 uri")

        bucket, object_name = _extract_minio_ref(material.uri)
        filename = _material_filename_from_object_name(object_name)
        extension = Path(filename).suffix.lower() if filename else None

        kind = _infer_upload_kind(mime_type=material.mime_type, extension=extension)

        st = storage or ObjectStorage()
        await st.ensure_buckets()
        content_bytes = await st.get_bytes(
            ObjectRef(bucket=bucket, object_name=object_name)
        )

        if kind in (None, "txt"):
            text = _decode_text_bytes(content_bytes)
            return _ensure_non_empty(
                ParsedDocument(
                    text=text,
                    mime_type=material.mime_type or "text/plain",
                    locator={"filename": filename} if filename else None,
                    metadata={"source_type": material.source_type.value},
                )
            )

        if kind == "md":
            text = _decode_text_bytes(content_bytes)
            return _ensure_non_empty(
                ParsedDocument(
                    text=text,
                    mime_type=material.mime_type or "text/markdown",
                    locator={"filename": filename} if filename else None,
                    metadata={"source_type": material.source_type.value},
                )
            )

        if kind == "docx":
            doc = await _parse_docx(content_bytes)
            doc.mime_type = material.mime_type or "text/markdown"
            doc.locator = {"filename": filename} if filename else None
            doc.metadata = {
                **(doc.metadata or {}),
                "source_type": material.source_type.value,
            }
            return _ensure_non_empty(doc)

        if kind == "pdf":
            mineru_error: ParseError | None = None
            try:
                doc = await _parse_pdf_with_mineru(content_bytes, settings=cfg)
                if not _has_meaningful_content(doc):
                    raise ParseError(
                        error_code="MINERU_EMPTY_RESULT", message="MinerU 解析结果为空"
                    )
            except ParseError as exc:
                mineru_error = exc
                fallback_enabled = bool(getattr(cfg, "pdf_fallback_enabled", True))
                if not fallback_enabled:
                    raise
                doc = await _parse_pdf_with_pypdf_fallback(
                    content_bytes,
                    settings=cfg,
                    mineru_error=exc,
                )

            doc.mime_type = material.mime_type or "application/pdf"
            doc.locator = {"filename": filename} if filename else None
            metadata = {
                **(doc.metadata or {}),
                "source_type": material.source_type.value,
            }
            if mineru_error is not None:
                metadata.setdefault("mineru_error_code", mineru_error.error_code)
                metadata.setdefault("mineru_error_message", mineru_error.message)
            doc.metadata = metadata
            return _ensure_non_empty(doc)

        raise ParseError(
            error_code="UNSUPPORTED_FILE_TYPE",
            message=f"不支持的上传文件类型：mime_type={material.mime_type!r} extension={extension!r}",
        )

    raise ParseError(
        error_code="UNSUPPORTED_SOURCE_TYPE",
        message=f"不支持的 source_type={material.source_type!r}",
    )


async def _parse_docx(content_bytes: bytes) -> ParsedDocument:
    try:
        from docx import Document  # type: ignore[import-not-found]
        from docx.oxml.table import CT_Tbl  # type: ignore[import-not-found]
        from docx.oxml.text.paragraph import CT_P  # type: ignore[import-not-found]
        from docx.table import Table  # type: ignore[import-not-found]
        from docx.text.paragraph import Paragraph  # type: ignore[import-not-found]
    except Exception as exc:  # pragma: no cover
        raise ParseError(
            error_code="PYTHON_DOCX_NOT_INSTALLED",
            message="未安装 python-docx，无法解析 .docx",
            details={"error": str(exc)},
        ) from exc

    def _render_table_md(table: Table) -> str:
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                txt = (cell.text or "").replace("\n", " ").strip()
                txt = txt.replace("|", "\\|")
                cells.append(txt)
            rows.append(cells)

        if not rows:
            return ""

        col_count = max(len(r) for r in rows) if rows else 0
        norm_rows = [r + [""] * (col_count - len(r)) for r in rows]
        header = norm_rows[0]
        body = norm_rows[1:] if len(norm_rows) > 1 else []

        md_lines = []
        md_lines.append("| " + " | ".join(header) + " |")
        md_lines.append("| " + " | ".join(["---"] * col_count) + " |")
        for r in body:
            md_lines.append("| " + " | ".join(r) + " |")
        return "\n".join(md_lines)

    def _iter_block_items(doc) -> list[Any]:
        blocks: list[Any] = []
        body = getattr(doc, "element", None)
        body = getattr(body, "body", None)
        if body is None:
            return blocks

        for child in body.iterchildren():
            if isinstance(child, CT_P):
                blocks.append(Paragraph(child, doc))
            elif isinstance(child, CT_Tbl):
                blocks.append(Table(child, doc))
        return blocks

    def _parse_sync() -> str:
        document = Document(BytesIO(content_bytes))
        lines: list[str] = []
        for block in _iter_block_items(document):
            if isinstance(block, Paragraph):
                text = (block.text or "").strip()
                if text:
                    lines.append(text)
                    lines.append("")
            elif isinstance(block, Table):
                table_md = _render_table_md(block).strip()
                if table_md:
                    lines.append(table_md)
                    lines.append("")
        return "\n".join(lines).strip()

    text = await asyncio.to_thread(_parse_sync)
    return ParsedDocument(
        text=text, mime_type="text/markdown", locator=None, metadata=None, chunks=None
    )


async def _parse_url(
    url: str, *, http_client: httpx.AsyncClient, settings: Settings
) -> ParsedDocument:
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"}:
        raise ParseError(
            error_code="INVALID_URL_SCHEME",
            message=f"不支持的 URL scheme: {parsed.scheme!r}",
        )

    # 这些配置项建议按需开放配置；默认值由 Settings 提供，缺失时再走兜底。
    max_redirects = getattr(settings, "ingestion_url_max_redirects", 3)
    max_bytes = getattr(settings, "ingestion_url_max_bytes", 20 * 1024 * 1024)
    user_agent = getattr(
        settings, "ingestion_url_user_agent", "multi-kb-agent/ingestion"
    )

    headers = {"User-Agent": user_agent}

    try:
        async with http_client.stream(
            "GET",
            url,
            follow_redirects=True,
            headers=headers,
        ) as resp:
            if len(resp.history) > max_redirects:
                raise ParseError(
                    error_code="URL_TOO_MANY_REDIRECTS",
                    message=f"URL 重定向次数过多（>{max_redirects}）",
                    details={"max_redirects": max_redirects, "url": url},
                )

            if resp.status_code >= 400:
                raise ParseError(
                    error_code="URL_FETCH_FAILED",
                    message=f"URL 抓取失败：HTTP {resp.status_code}",
                    details={"status_code": resp.status_code, "url": url},
                )

            buf = bytearray()
            async for chunk in resp.aiter_bytes():
                buf.extend(chunk)
                if len(buf) > max_bytes:
                    raise ParseError(
                        error_code="URL_RESPONSE_TOO_LARGE",
                        message=f"URL 响应体超过限制（>{max_bytes} bytes）",
                        details={"max_bytes": max_bytes, "url": url},
                    )

            content_bytes = bytes(buf)
    except ParseError:
        raise
    except httpx.TooManyRedirects as exc:
        raise ParseError(
            error_code="URL_TOO_MANY_REDIRECTS",
            message=f"URL 重定向次数过多（>{max_redirects}）",
            details={"max_redirects": max_redirects, "url": url},
        ) from exc
    except Exception as exc:
        raise ParseError(
            error_code="URL_FETCH_EXCEPTION",
            message=f"URL 抓取异常：{exc}",
            details={"url": url},
        ) from exc

    # 编码：尽量利用 httpx 推断；否则回退到 utf-8 或 gb18030。
    try:
        html_text = content_bytes.decode("utf-8")
    except UnicodeDecodeError:
        html_text = content_bytes.decode("gb18030", errors="replace")

    try:
        from readability import Document  # type: ignore[import-not-found]
    except Exception as exc:  # pragma: no cover
        raise ParseError(
            error_code="READABILITY_NOT_INSTALLED",
            message="未安装 readability-lxml，无法进行 URL 正文抽取",
            details={"error": str(exc)},
        ) from exc

    readability_doc = Document(html_text)
    title = (readability_doc.title() or "").strip()
    main_html = readability_doc.summary()

    # Readability 输出通常为 HTML，这里转为纯文本；纯文本本身也可作为合法 Markdown。
    try:
        from lxml import html as lxml_html  # type: ignore[import-not-found]

        root = lxml_html.fromstring(main_html)
        text = (root.text_content() or "").strip()
    except Exception as exc:
        raise ParseError(
            error_code="URL_EXTRACT_FAILED",
            message=f"URL 正文抽取失败：{exc}",
            details={"url": url},
        ) from exc

    # 规范化空行，避免大量空白影响分块与 embedding。
    lines = [ln.strip() for ln in text.splitlines()]
    lines = [ln for ln in lines if ln]
    body_text = "\n".join(lines).strip()

    if title:
        md = f"# {title}\n\n{body_text}".strip()
    else:
        md = body_text

    return ParsedDocument(
        text=md,
        mime_type="text/markdown",
        locator={"url": url},
        metadata={"title": title} if title else None,
        chunks=None,
    )


def _resolve_mineru_options(settings: Settings) -> tuple[str, str, bool, bool]:
    lang = str(getattr(settings, "mineru_lang", "ch") or "ch").strip() or "ch"

    parse_method = (
        str(getattr(settings, "mineru_parse_method", "auto") or "auto").strip().lower()
    )
    if parse_method not in {"auto", "txt", "ocr"}:
        parse_method = "auto"

    formula_enable = bool(getattr(settings, "mineru_formula_enable", True))
    table_enable = bool(getattr(settings, "mineru_table_enable", True))
    return lang, parse_method, formula_enable, table_enable


async def _parse_pdf_with_mineru(
    content_bytes: bytes, *, settings: Settings
) -> ParsedDocument:
    """使用 MinerU Python API 解析 PDF，返回 block=chunk 的 chunks 与全文。"""
    lang, parse_method, formula_enable, table_enable = _resolve_mineru_options(settings)

    def _parse_sync() -> dict[str, Any]:
        try:
            from mineru.backend.pipeline.model_json_to_middle_json import (  # type: ignore[import-not-found]
                result_to_middle_json,
            )
            from mineru.backend.pipeline.pipeline_analyze import (  # type: ignore[import-not-found]
                doc_analyze,
            )
            from mineru.cli.common import (  # type: ignore[import-not-found]
                convert_pdf_bytes_to_bytes_by_pypdfium2,
            )
            from mineru.data.data_reader_writer import FileBasedDataWriter  # type: ignore[import-not-found]
        except Exception as exc:
            raise ParseError(
                error_code="MINERU_NOT_INSTALLED",
                message="未安装 MinerU（mineru[core]），无法解析 PDF",
                details={"error": str(exc)},
            ) from exc

        model_source = getattr(settings, "mineru_model_source", None)
        if isinstance(model_source, str) and model_source.strip():
            os.environ["MINERU_MODEL_SOURCE"] = model_source.strip()

        try:
            pdf_bytes = convert_pdf_bytes_to_bytes_by_pypdfium2(content_bytes, 0, None)
        except TimeoutError as exc:
            raise ParseError(
                error_code="MINERU_TIMEOUT",
                message="MinerU PDF 预处理超时",
                details={"parse_method": parse_method, "lang": lang},
            ) from exc
        except Exception as exc:
            raise ParseError(
                error_code="MINERU_RUNTIME_ERROR",
                message=f"MinerU PDF 预处理失败：{exc}",
                details={"parse_method": parse_method, "lang": lang},
            ) from exc

        try:
            (
                infer_results,
                all_image_lists,
                all_pdf_docs,
                lang_list,
                ocr_enabled_list,
            ) = doc_analyze(
                [pdf_bytes],
                [lang],
                parse_method=parse_method,
                formula_enable=formula_enable,
                table_enable=table_enable,
            )
        except TimeoutError as exc:
            raise ParseError(
                error_code="MINERU_TIMEOUT",
                message="MinerU 解析超时",
                details={"parse_method": parse_method, "lang": lang},
            ) from exc
        except Exception as exc:
            raise ParseError(
                error_code="MINERU_RUNTIME_ERROR",
                message=f"MinerU 解析失败：{exc}",
                details={"parse_method": parse_method, "lang": lang},
            ) from exc

        try:
            model_list = infer_results[0]
            images_list = all_image_lists[0]
            pdf_doc = all_pdf_docs[0]
            parsed_lang = lang_list[0]
            ocr_enabled = ocr_enabled_list[0]
        except Exception as exc:
            raise ParseError(
                error_code="MINERU_BAD_OUTPUT",
                message="MinerU 输出结构异常",
                details={"parse_method": parse_method, "lang": lang},
            ) from exc

        try:
            with tempfile.TemporaryDirectory(prefix="mineru_") as tmp:
                image_writer = FileBasedDataWriter(tmp)
                middle_json = result_to_middle_json(
                    model_list,
                    images_list,
                    pdf_doc,
                    image_writer,
                    parsed_lang,
                    ocr_enabled,
                    formula_enable,
                )
        except TimeoutError as exc:
            raise ParseError(
                error_code="MINERU_TIMEOUT",
                message="MinerU 生成 middle.json 超时",
                details={"parse_method": parse_method, "lang": lang},
            ) from exc
        except ParseError:
            raise
        except Exception as exc:
            raise ParseError(
                error_code="MINERU_RUNTIME_ERROR",
                message=f"MinerU 生成 middle.json 失败：{exc}",
                details={"parse_method": parse_method, "lang": lang},
            ) from exc

        if not isinstance(middle_json, dict):
            raise ParseError(
                error_code="MINERU_BAD_OUTPUT",
                message="MinerU 输出 middle_json 非 dict",
            )
        return middle_json

    middle_json = await asyncio.to_thread(_parse_sync)
    chunks = extract_pdf_chunks_from_middle_json(middle_json)
    full_text = "\n\n".join([c.text for c in chunks]).strip()
    metadata: dict[str, Any] = {
        "pdf_parse_path": "mineru",
        "fallback_used": False,
        "mineru_block_count": len(chunks),
        "mineru_backend": middle_json.get("_backend"),
        "mineru_version": middle_json.get("_version_name"),
        "mineru_lang": lang,
        "mineru_parse_method": parse_method,
        "mineru_formula_enable": formula_enable,
        "mineru_table_enable": table_enable,
    }
    return ParsedDocument(
        text=full_text,
        mime_type="application/pdf",
        locator=None,
        metadata=metadata,
        chunks=chunks,
    )


async def _parse_pdf_with_pypdf_fallback(
    content_bytes: bytes,
    *,
    settings: Settings,
    mineru_error: ParseError | None = None,
) -> ParsedDocument:
    max_pages = max(int(getattr(settings, "pdf_fallback_max_pages", 500)), 1)
    min_chars = max(int(getattr(settings, "pdf_fallback_min_text_chars", 20)), 0)

    def _parse_sync() -> tuple[str, int]:
        try:
            from pypdf import PdfReader  # type: ignore[import-not-found]
        except Exception as exc:  # pragma: no cover
            raise ParseError(
                error_code="PYPDF_NOT_INSTALLED",
                message="未安装 pypdf，无法执行 PDF 文本兜底",
                details={"error": str(exc)},
            ) from exc

        try:
            reader = PdfReader(BytesIO(content_bytes))
        except Exception as exc:
            raise ParseError(
                error_code="PDF_FALLBACK_FAILED",
                message=f"pypdf 打开 PDF 失败：{exc}",
            ) from exc

        page_total = len(reader.pages)
        pages_to_read = min(page_total, max_pages)
        page_texts: list[str] = []
        for page_index in range(pages_to_read):
            try:
                raw = reader.pages[page_index].extract_text() or ""
            except Exception:
                raw = ""
            text = raw.strip()
            if text:
                page_texts.append(text)
        return "\n\n".join(page_texts).strip(), pages_to_read

    text, pages_to_read = await asyncio.to_thread(_parse_sync)
    if len(text) < min_chars:
        details: dict[str, Any] = {
            "min_text_chars": min_chars,
            "extracted_chars": len(text),
            "pages_processed": pages_to_read,
        }
        if mineru_error is not None:
            details["mineru_error_code"] = mineru_error.error_code
        raise ParseError(
            error_code="PDF_FALLBACK_EMPTY",
            message="PDF 文本兜底结果过短或为空",
            details=details,
        )

    metadata: dict[str, Any] = {
        "pdf_parse_path": "fallback_pypdf",
        "fallback_used": True,
        "mineru_block_count": 0,
        "fallback_pages_processed": pages_to_read,
        "fallback_max_pages": max_pages,
        "fallback_min_text_chars": min_chars,
    }
    if mineru_error is not None:
        metadata["mineru_error_code"] = mineru_error.error_code

    return ParsedDocument(
        text=text,
        mime_type="application/pdf",
        locator=None,
        metadata=metadata,
        chunks=None,
    )
